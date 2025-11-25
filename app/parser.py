import re
import sys
from typing import Dict, Optional

import pdfplumber
from docx import Document

# 设置标准输出编码为 UTF-8，避免 Windows 控制台编码问题
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')


def extract_text_from_pdf(pdf_path: str) -> str:
    """从 PDF 中提取全部文字。"""
    text_parts = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                # 安全处理文本，移除可能导致编码问题的字符
                if t:
                    # 确保文本是有效的 UTF-8
                    try:
                        t = t.encode('utf-8', errors='replace').decode('utf-8')
                    except:
                        t = ""
                text_parts.append(t)
    except Exception as e:
        # 如果提取失败，返回空字符串而不是抛出异常
        error_msg = str(e).encode('utf-8', errors='replace').decode('utf-8')
        raise ValueError(f"PDF 文本提取失败: {error_msg}")
    return "\n".join(text_parts)


def extract_text_from_docx(docx_path: str) -> str:
    """
    从 DOCX 文件尽量提取全部文本：
    - 普通段落
    - 表格中的文本
    注意：页眉/页脚/文本框等复杂对象可能仍然获取不到。
    """
    parts = []
    try:
        doc = Document(docx_path)
        
        # 1) 表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text_cells = []
                for cell in row.cells:
                    cell_text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                    if cell_text:
                        # 安全处理文本
                        try:
                            cell_text = cell_text.encode('utf-8', errors='replace').decode('utf-8')
                        except:
                            continue
                        row_text_cells.append(cell_text)
                if row_text_cells:
                    parts.append(" | ".join(row_text_cells))
        
        # 2) 普通段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # 安全处理文本
                try:
                    text = text.encode('utf-8', errors='replace').decode('utf-8')
                except:
                    continue
                parts.append(text)
    except Exception as e:
        # 如果提取失败，返回错误信息
        error_msg = str(e).encode('utf-8', errors='replace').decode('utf-8')
        raise ValueError(f"DOCX 文本提取失败: {error_msg}")
    
    # 如需更多结构（页眉、页脚等），需要用更底层的 API，这里先不展开
    return "\n".join(parts)


def extract_docx_to_html(docx_path: str) -> str:
    """
    将 DOCX 文件内容转换为 HTML 格式，用于预览
    """
    try:
        doc = Document(docx_path)
        html_parts = ['<div class="docx-preview">']
        
        # 处理段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # 安全处理文本
                try:
                    text = text.encode('utf-8', errors='replace').decode('utf-8')
                    # 转义 HTML 特殊字符
                    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    # 检测是否为标题（简单判断：短文本且可能包含冒号）
                    if len(text) < 50 and '：' in text:
                        html_parts.append(f'<h3>{text}</h3>')
                    else:
                        html_parts.append(f'<p>{text}</p>')
                except:
                    continue
        
        # 处理表格
        for table in doc.tables:
            html_parts.append('<table class="docx-table">')
            for row in table.rows:
                html_parts.append('<tr>')
                for cell in row.cells:
                    cell_text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                    if cell_text:
                        try:
                            cell_text = cell_text.encode('utf-8', errors='replace').decode('utf-8')
                            cell_text = cell_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            html_parts.append(f'<td>{cell_text}</td>')
                        except:
                            html_parts.append('<td></td>')
                    else:
                        html_parts.append('<td></td>')
                html_parts.append('</tr>')
            html_parts.append('</table>')
        
        html_parts.append('</div>')
        return '\n'.join(html_parts)
    except Exception as e:
        error_msg = str(e).encode('utf-8', errors='replace').decode('utf-8')
        raise ValueError(f"DOCX 转 HTML 失败: {error_msg}")


def parse_candidate_info(text: str) -> Dict[str, Optional[str]]:
    """
    通用文本解析（备用）。
    简单解析：邮箱、电话、学校、学历、专业。
    """

    lines = [line.strip() for line in text.splitlines() if line.strip()]

    email = None
    email_pattern = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
    for line in lines:
        m = email_pattern.search(line)
        if m:
            email = m.group(0)
            break

    phone = None
    phone_pattern = re.compile(r"1[3-9]\d{9}")
    for line in lines:
        m = phone_pattern.search(line)
        if m:
            phone = m.group(0)
            break

    name = None
    if lines:
        m_name = re.match(r"([\u4e00-\u9fa5]{2,4})", lines[0])
        if m_name:
            name = m_name.group(1)

    university = None
    for line in lines:
        if "大学" in line or "学院" in line:
            m = re.search(r"(.+(大学|学院))", line)
            if m:
                university = m.group(1).strip()
                break

    degree = None
    for d in ["博士", "硕士", "本科", "大专", "中专"]:
        if any(d in line for line in lines):
            degree = d
            break

    major = None
    for line in lines:
        if "专业" in line:
            m = re.search(r"专业[:：\s]*([\u4e00-\u9fa5A-Za-z0-9（）()·\s]{2,30})", line)
            if m:
                major = m.group(1).strip()
                break

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "university": university,
        "degree": degree,
        "major": major,
    }


def parse_resume_text_cn(text: str) -> Dict[str, Optional[str]]:
    # 清理文本：移除不可打印字符和可能导致编码问题的字符
    text = text.replace("\u00A0", " ").strip()
    # 移除控制字符，但保留换行符和制表符
    text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t\r')
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    name = None
    gender = None
    age = None
    degree = None
    university = None
    major = None
    phone = None
    email = None

    if not lines:
        return {
            "name": None,
            "gender": None,
            "age": None,
            "phone": None,
            "email": None,
            "university": None,
            "degree": None,
            "major": None,
        }

    # 1) 姓名：第一行
    first_line = lines[0]
    m_name = re.match(r"^([\u4e00-\u9fa5]{2,4})$", first_line)
    if m_name:
        name = m_name.group(1)

    # 2) 性别 / 年龄 / 学历（前几行合并）
    base_info_lines = lines[1:4]
    base_info_text = " ".join(base_info_lines)

    if "女" in base_info_text:
        gender = "女"
    elif "男" in base_info_text:
        gender = "男"

    m_age = re.search(r"(\d{1,2})\s*岁", base_info_text)
    if m_age:
        age = m_age.group(1)

    for d in ["博士", "硕士", "本科", "大专", "中专"]:
        if d in base_info_text:
            degree = d
            break

    # 3) 联系方式
    phone_pattern = re.compile(r"1[3-9]\d{9}")
    m_phone = phone_pattern.search(text)
    if m_phone:
        phone = m_phone.group(0)

    email_pattern = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
    m_email = email_pattern.search(text)
    if m_email:
        email = m_email.group(0)

    # 4) 教育经历
    edu_start_idx = -1
    for i, line in enumerate(lines):
        if "教育经历" in line:
            edu_start_idx = i
            break

    if edu_start_idx != -1:
        edu_lines = lines[edu_start_idx + 1: edu_start_idx + 6]
        edu_text = " ".join(edu_lines)

        # 学校
        for line in edu_lines:
            if "大学" in line or "学院" in line:
                m_uni = re.search(r"(.+?(大学|学院))", line)
                if m_uni:
                    university = m_uni.group(1).strip()
                else:
                    university = line.strip()
                break

        # 学历（以教育经历为准再覆盖）
        for d in ["博士", "硕士", "本科", "大专", "中专"]:
            if d in edu_text:
                degree = d
                break

        # 专业：针对你这行的格式精细拆分
        # 示例行：“本科学历 丨 学士学位 丨 康复治疗学（中外合作 丨 全日制”
        for line in edu_lines:
            if "学历" in line or "学位" in line:
                # 用多种分隔符拆分：丨、|、空格
                # 先统一替换为同一种分隔符，方便 split
                tmp = line.replace("|", "丨")
                # 以“丨”切片
                parts = [p.strip() for p in tmp.split("丨") if p.strip()]
                # 过滤掉明显不是专业的片段：包含“学历”、“学位”、“全日制”等
                candidates = []
                for p in parts:
                    if any(key in p for key in ["学历", "学位", "全日制", "统招", "非全日制"]):
                        continue
                    # 要有至少2个中文/字母，避免一些杂质词
                    if len(re.findall(r"[\u4e00-\u9fa5A-Za-z]", p)) >= 2:
                        candidates.append(p)

                # 在候选中，选择第一个作为专业
                if candidates:
                    major = candidates[0]
                break

    return {
        "name": name,
        "gender": gender,
        "age": age,
        "phone": phone,
        "email": email,
        "university": university,
        "degree": degree,
        "major": major,
    }

def parse_resume_file(path: str) -> Dict[str, Optional[str]]:
    """
    统一入口：根据文件扩展名选择解析方式（PDF / DOCX），
    优先使用中文简历解析规则。
    """
    try:
        lower = path.lower()
        if lower.endswith(".pdf"):
            text = extract_text_from_pdf(path)
            return parse_resume_text_cn(text)
        elif lower.endswith(".docx"):
            text = extract_text_from_docx(path)
            return parse_resume_text_cn(text)
        else:
            raise ValueError("暂不支持的简历格式（仅支持 PDF / DOCX）")
    except Exception as e:
        # 安全处理异常，避免编码问题
        error_msg = str(e)
        try:
            error_msg = error_msg.encode('utf-8', errors='replace').decode('utf-8')
        except:
            error_msg = "解析文件时发生未知错误"
        raise ValueError(error_msg) from e
